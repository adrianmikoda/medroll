import os
import json
import fitz
from docx import Document
from pydantic import BaseModel
from langdetect import detect_langs
from abc import ABC, abstractmethod

#Pydantic
class MedicalDocument(BaseModel):
    session_id: str
    filename: str
    content: str
    language: str
    status: str = "processed"

class BaseExtractor(ABC):
    @abstractmethod
    def extract(self, file_path: str) -> str:
        pass

class PDFExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text.strip()

class DocxExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs]).strip()

class FileConverter:
    EXTRACTORS = {
        '.pdf': PDFExtractor(),
        '.docx': DocxExtractor(),
    }

    def __init__(self, file_path: str, session_id: str):
        self.file_path = file_path
        self.session_id = session_id
        self.extension = os.path.splitext(file_path)[1].lower()

    def detect_language(self, text: str) -> tuple[bool, str]:
        try:
            predictions = detect_langs(text) # [lang:confidence score]
            top_lang = predictions[0]

            if top_lang.lang not in ['pl', 'de', 'en']:
                return False, f"The language {top_lang.lang} is not supported"
            
            if top_lang.prob < 0.8:
                return False, "Validation Error: Multiple Languages detected. Only single-language documents are supported"

            return True, top_lang.lang
        except:
            return False, "Could not recognize the language"

    def convert_to_json(self) -> str:
        try:
            extractor = self.EXTRACTORS.get(self.extension)
            if not extractor:
                raise ValueError(f"Unsupported format: {self.extension}")

            raw_text = extractor.extract(self.file_path)

            is_valid_lang, lang_result = self.detect_language(raw_text)
            if not is_valid_lang:
                raise ValueError(lang_result)
            
            doc_obj = MedicalDocument(
                session_id=self.session_id,
                filename=os.path.basename(self.file_path),
                content = raw_text,
                language = lang_result
            )

            return doc_obj.model_dump_json(indent=4)
        except Exception as e:
            # API Errors
            return json.dumps({"status": "error", "message": str(e)}, ensure_ascii=False)